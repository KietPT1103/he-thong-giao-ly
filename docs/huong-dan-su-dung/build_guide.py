from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSETS = Path(__file__).resolve().parent / "assets"
OUTPUT = ROOT / "docs" / "Huong-dan-su-dung-he-thong-giao-ly.docx"
LOGO = ROOT / "public" / "images" / "02_individual_assets" / "logo-full.png"

# compact_reference_guide preset + one named editorial-cover override.
FONT = "Calibri"
NAVY = "0B214D"
BLUE = "2563EB"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
TEXT = "1E293B"
MUTED = "64748B"
BORDER = "D9E2EE"
PALE_BLUE = "F4F8FF"
TABLE_HEADER = "E8EEF5"
PALE_GREEN = "EDF8F0"
PALE_AMBER = "FFF8E8"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, *, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 font: str = FONT) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_fixed_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}, got {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, HEADING_BLUE, 18, 10),
        "Heading 2": (13, HEADING_BLUE, 14, 7),
        "Heading 3": (12, HEADING_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_together = True


def add_numbering_definition(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.append(size)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int, *, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        set_run_font(lead, size=11, color=TEXT, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11, color=TEXT)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=TEXT)
    return paragraph


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, color=TEXT, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=TEXT)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=TEXT)
    return paragraph


def add_callout(doc: Document, label: str, text: str, *, tone: str = "blue") -> None:
    fill = {"blue": PALE_BLUE, "green": PALE_GREEN, "amber": PALE_AMBER}[tone]
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), "A9CCFF" if tone == "blue" else BORDER)
        borders.append(node)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=NAVY, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=TEXT)


def add_screenshot(doc: Document, filename: str, caption_text: str, alt_text: str) -> None:
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.35))
    doc_pr = run._element.find(".//" + qn("wp:docPr"))
    if doc_pr is not None:
        doc_pr.set("descr", alt_text)
        doc_pr.set("title", caption_text)
    caption = doc.add_paragraph(caption_text, style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_role_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    widths = [1728, 3024, 4608]
    headers = ["Vai trò", "Trách nhiệm chính", "Khu vực sử dụng thường xuyên"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_HEADER)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=10, color=NAVY, bold=True)
    rows = [
        ("Quản trị viên", "Thiết lập dữ liệu nền và phân quyền", "Tài khoản, giáo xứ, giáo lý viên, phụ huynh, thiếu nhi, lớp học"),
        ("Giáo lý viên", "Vận hành lớp được phân công", "Lịch dạy, điểm danh, bài tập, chấm bài, thông báo lớp"),
        ("Thiếu nhi", "Tham gia học tập và điểm danh", "Bài tập, thông báo, QR điểm danh, lịch học"),
        ("Phụ huynh", "Theo dõi hồ sơ và hoạt động của con", "Các con của tôi, lịch học, tham dự, bài tập, thông báo"),
    ]
    for role, responsibility, modules in rows:
        cells = table.add_row().cells
        for index, text in enumerate((role, responsibility, modules)):
            paragraph = cells[index].paragraphs[0]
            run = paragraph.add_run(text)
            set_run_font(run, size=10, color=TEXT, bold=index == 0)
    repeat_table_header(table.rows[0])
    set_fixed_table_geometry(table, widths)
    set_table_borders(table)
    doc.add_paragraph()


def add_demo_accounts_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    widths = [1800, 4560, 3000]
    for index, text in enumerate(("Vai trò", "Tài khoản mẫu", "Mật khẩu local")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_HEADER)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=10, color=NAVY, bold=True)
    for role, email in (
        ("Quản trị viên", "admin@giaoly.test"),
        ("Giáo lý viên", "teacher@giaoly.test"),
        ("Thiếu nhi", "child@giaoly.test"),
        ("Phụ huynh", "parent@giaoly.test"),
    ):
        cells = table.add_row().cells
        for index, text in enumerate((role, email, "password")):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            set_run_font(run, size=10, color=TEXT, bold=index == 0)
    repeat_table_header(table.rows[0])
    set_fixed_table_geometry(table, widths)
    set_table_borders(table)


def set_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Hành Trang Đức Tin  |  Hướng dẫn sử dụng")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=8.5, color=MUTED)
    field_begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    field_begin_run._r.append(fld_begin)

    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instruction_run._r.append(instr)

    separator_run = paragraph.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separator_run._r.append(fld_separate)

    page_run = paragraph.add_run("1")
    set_run_font(page_run, size=8.5, color=MUTED)

    field_end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_end_run._r.append(fld_end)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(28)
    logo_run = logo_p.add_run()
    logo_run.add_picture(str(LOGO), width=Inches(2.45))
    logo_pr = logo_run._element.find(".//" + qn("wp:docPr"))
    if logo_pr is not None:
        logo_pr.set("descr", "Logo hệ thống Hành Trang Đức Tin - Giáo xứ Cái Răng.")
        logo_pr.set("title", "Hành Trang Đức Tin")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("SỔ TAY VẬN HÀNH HỆ THỐNG")
    set_run_font(run, size=10.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Hướng dẫn sử dụng\nHành Trang Đức Tin")
    set_run_font(run, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("Dành cho Quản trị viên, Giáo lý viên, Thiếu nhi và Phụ huynh")
    set_run_font(run, size=13, color=MUTED)

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(8)
    run = metadata.add_run("Phiên bản 1.0  |  Cập nhật ngày 02/09/2026")
    set_run_font(run, size=10.5, color=TEXT, bold=True)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_after = Pt(0)
    run = scope.add_run("Giáo xứ Cái Răng - Giáo phận Cần Thơ")
    set_run_font(run, size=10.5, color=MUTED, italic=True)
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    doc.core_properties.title = "Hướng dẫn sử dụng hệ thống Hành Trang Đức Tin"
    doc.core_properties.subject = "Sổ tay sử dụng theo vai trò"
    doc.core_properties.author = "Ban Giáo lý Giáo xứ Cái Răng"
    doc.core_properties.keywords = "giáo lý, quản trị, điểm danh, bài tập, hướng dẫn"
    doc.core_properties.comments = "Tài liệu được tạo từ giao diện hệ thống local đã kiểm tra."

    configure_styles(doc)
    number_list = add_numbering_definition(doc, bullet=False)
    bullet_list = add_numbering_definition(doc, bullet=True)
    set_running_header_footer(doc)
    add_cover(doc)

    doc.add_heading("Giới thiệu", level=1)
    add_body(doc, "Tài liệu này hướng dẫn các thao tác đang hoạt động trên hệ thống Hành Trang Đức Tin. Mỗi phần được sắp xếp theo vai trò đăng nhập và đi kèm ảnh màn hình thực tế để người dùng dễ đối chiếu.")
    add_callout(doc, "Phạm vi", "Tên menu có thể thay đổi theo quyền được cấp. Nếu một chức năng không xuất hiện, hãy liên hệ quản trị viên để kiểm tra vai trò và quyền truy cập.")
    doc.add_heading("Vai trò trong hệ thống", level=2)
    add_role_table(doc)
    doc.add_heading("Quy ước thao tác", level=2)
    for text in (
        "Nút màu xanh là hành động chính như Lưu, Cập nhật, Phát hành hoặc Điểm danh.",
        "Dấu * cạnh nhãn cho biết trường bắt buộc; hệ thống sẽ báo lỗi ngay tại trường nếu dữ liệu chưa hợp lệ.",
        "Biểu tượng chuông ở góc trên mở thông báo; khu vực tài khoản ở cuối thanh bên dùng để mở hồ sơ hoặc đăng xuất.",
        "Trên điện thoại, thanh điều hướng được thu gọn; sử dụng nút menu ở góc trên để mở danh sách chức năng.",
    ):
        add_list_item(doc, text, bullet_list)

    doc.add_page_break()
    doc.add_heading("1. Đăng nhập và sử dụng chung", level=1)
    doc.add_heading("1.1 Đăng nhập", level=2)
    for text in (
        "Mở địa chỉ hệ thống do Ban Giáo lý cung cấp.",
        "Nhập Email và Mật khẩu của tài khoản.",
        "Chọn Đăng nhập. Hệ thống tự chuyển đến không gian đúng với vai trò của tài khoản.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "01-dang-nhap.png", "Hình 1 - Màn hình đăng nhập hệ thống.", "Màn hình đăng nhập gồm trường email, mật khẩu và nút Đăng nhập.")
    add_callout(doc, "Bảo mật", "Không dùng chung tài khoản. Sau khi sử dụng trên thiết bị công cộng, mở menu tài khoản và chọn Đăng xuất.", tone="amber")

    doc.add_heading("1.2 Đăng ký tài khoản", level=2)
    add_body(doc, "Người dùng chưa có tài khoản có thể tự đăng ký với vai trò Thiếu nhi hoặc Phụ huynh. Tài khoản Quản trị viên và Giáo lý viên do quản trị viên hệ thống cấp, không đăng ký tại màn hình này.")
    for text in (
        "Tại màn hình Đăng nhập, chọn Đăng ký ngay hoặc mở đường dẫn /register.",
        "Nhập Họ và tên, Email và Số điện thoại Việt Nam đang sử dụng.",
        "Chọn đúng vai trò Thiếu nhi hoặc Phụ huynh. Vai trò quyết định không gian được mở sau khi đăng ký.",
        "Nhập Mật khẩu có ít nhất 8 ký tự, sau đó nhập lại chính xác ở trường Xác nhận mật khẩu.",
        "Chọn Tạo tài khoản. Khi đăng ký thành công, hệ thống tự đăng nhập và chuyển đến không gian của vai trò đã chọn.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "01b-dang-ky.png", "Hình 2 - Màn hình đăng ký tài khoản.", "Biểu mẫu đăng ký gồm thông tin cá nhân, lựa chọn vai trò Thiếu nhi hoặc Phụ huynh và mật khẩu xác nhận.")
    add_callout(doc, "Sau khi đăng ký", "Tài khoản Thiếu nhi mới chưa tự được xếp lớp; tài khoản Phụ huynh mới chưa tự liên kết với hồ sơ con. Hãy liên hệ quản trị viên để hoàn tất xếp lớp hoặc liên kết gia đình.", tone="green")

    doc.add_heading("1.3 Cập nhật tài khoản", level=2)
    add_body(doc, "Chọn khu vực tài khoản ở cuối thanh bên, sau đó chọn Tài khoản. Tại đây người dùng có thể xem và cập nhật thông tin cá nhân theo quyền được cấp.")

    doc.add_page_break()
    doc.add_heading("2. Hướng dẫn dành cho Quản trị viên", level=1)
    add_body(doc, "Quản trị viên thiết lập dữ liệu nền, quản lý tài khoản và tổ chức lớp học. Chỉ thay đổi dữ liệu khi đã xác nhận đúng giáo xứ, niên khóa và đối tượng liên quan.")
    doc.add_heading("2.1 Theo dõi tổng quan", level=2)
    for text in (
        "Mở Tổng quan để kiểm tra số lượng giáo xứ, giáo lý viên, thiếu nhi và lớp đang hoạt động.",
        "Theo dõi các chỉ số chuyên cần và phiên điểm danh gần đây để phát hiện dữ liệu cần xử lý.",
    ):
        add_list_item(doc, text, bullet_list)
    add_screenshot(doc, "02-admin-tong-quan.png", "Hình 3 - Không gian tổng quan của Quản trị viên.", "Trang tổng quan quản trị với các chỉ số vận hành hệ thống.")

    doc.add_heading("2.2 Quản lý dữ liệu hệ thống", level=2)
    add_body(doc, "Thanh bên cho phép truy cập Quản lý tài khoản, Giáo xứ, Giáo lý viên, Phụ huynh, Thiếu nhi, Danh mục lớp học, Lớp học và Thông báo. Trong mỗi danh sách, sử dụng ô tìm kiếm và bộ lọc trước khi tạo mới hoặc cập nhật.")
    add_screenshot(doc, "03-admin-danh-sach-lop.png", "Hình 4 - Danh sách lớp học và bộ lọc quản trị.", "Danh sách lớp học với bộ lọc giáo xứ, niên khóa, khối và trạng thái.")

    doc.add_heading("2.3 Thiết lập Danh mục lớp học", level=2)
    add_body(doc, "Mở Danh mục lớp học để quản lý dữ liệu dùng chung cho các lớp trong từng giáo xứ. Màn hình gồm ba tab Niên khóa, Khối giáo lý và Phòng học.")
    for text in (
        "Chọn đúng Giáo xứ đang quản lý trước khi thao tác.",
        "Ở tab Niên khóa, nhập tên, ngày bắt đầu, ngày kết thúc và chọn Niên khóa hiện tại khi cần. Mỗi giáo xứ chỉ có một niên khóa hiện tại.",
        "Ở tab Khối giáo lý, nhập tên khối, mã khối và thứ tự hiển thị.",
        "Ở tab Phòng học, nhập tên phòng và sức chứa; sức chứa có thể để trống nếu chưa xác định.",
        "Dùng nút bút chì để chỉnh sửa; dùng nút nguồn để Ngừng sử dụng hoặc Sử dụng lại danh mục.",
        "Chỉ xóa danh mục chưa có lớp liên kết. Nếu nút xóa bị khóa, hãy chọn Ngừng sử dụng để giữ an toàn dữ liệu lịch sử.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "04-admin-danh-muc-lop.png", "Hình 5 - Màn hình Danh mục lớp học của Quản trị viên.", "Màn hình quản lý Niên khóa, Khối giáo lý và Phòng học theo giáo xứ, kèm trạng thái và số lớp sử dụng.")
    add_callout(doc, "Quan trọng", "Tên danh mục là dữ liệu dùng chung. Khi đổi tên, mọi lớp đang sử dụng danh mục đó sẽ hiển thị tên mới; vì vậy hãy kiểm tra số lớp sử dụng trước khi lưu.", tone="amber")

    doc.add_heading("2.4 Chỉnh sửa và tổ chức lớp học", level=2)
    for text in (
        "Từ Danh sách lớp học, chọn dòng lớp cần chỉnh sửa.",
        "Kiểm tra Tên lớp, Mã lớp, Giáo xứ và Trạng thái. Với Niên khóa, Khối giáo lý và Phòng học, chỉ chọn giá trị từ danh mục đã thiết lập; không đổi tên trực tiếp tại biểu mẫu lớp.",
        "Nếu cần bổ sung hoặc đổi tên dữ liệu dùng chung, chọn Quản lý danh mục ở đầu biểu mẫu, cập nhật danh mục rồi quay lại lớp.",
        "Chọn Cập nhật lớp học để lưu; chọn Hủy bỏ để trả biểu mẫu về dữ liệu đã lưu gần nhất.",
        "Trong Tổng quan lớp học, dùng Quản lý danh sách, Phân công và Thiết lập để quản lý thiếu nhi, giáo lý viên và lịch học.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "04-admin-chinh-sua-lop.png", "Hình 6 - Màn hình chỉnh sửa lớp học.", "Màn hình chỉnh sửa lớp với nút Quản lý danh mục và các trường chọn Niên khóa, Khối giáo lý, Phòng học.")
    add_callout(doc, "Lưu ý", "Giáo lý viên không tự tạo lịch dạy. Quản trị viên thiết lập lịch học và phân công giáo lý viên phụ trách lớp.")

    doc.add_page_break()
    doc.add_heading("3. Hướng dẫn dành cho Giáo lý viên", level=1)
    doc.add_heading("3.1 Tổng quan và lớp phụ trách", level=2)
    add_body(doc, "Trang Tổng quan giúp giáo lý viên theo dõi lớp phụ trách, số thiếu nhi, phiên điểm danh và các công việc cần xử lý.")
    add_screenshot(doc, "05-giao-ly-vien-tong-quan.png", "Hình 7 - Tổng quan của Giáo lý viên.", "Trang tổng quan giáo lý viên với số liệu và lớp phụ trách.")
    add_body(doc, "Mở Lớp của tôi, sau đó chọn một lớp để xem danh sách thiếu nhi, lịch học định kỳ và giáo lý viên phụ trách.")
    add_screenshot(doc, "06-giao-ly-vien-chi-tiet-lop.png", "Hình 8 - Chi tiết lớp học của Giáo lý viên.", "Trang chi tiết lớp học với thông tin lớp, danh sách thiếu nhi và lịch học.")

    doc.add_heading("3.2 Điểm danh và xem lịch sử phiên", level=2)
    for text in (
        "Mở Điểm danh lớp và chọn lớp được phân công.",
        "Chọn phiên đang diễn ra hoặc mở phiên mới nếu chưa có phiên phù hợp.",
        "Đánh dấu trạng thái từng thiếu nhi và lưu điểm danh.",
        "Mở tab Danh sách phiên để lọc các phiên đang diễn ra, đã kết thúc hoặc đã hủy.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "07-danh-sach-phien-diem-danh.png", "Hình 9 - Danh sách phiên điểm danh của lớp.", "Danh sách phiên điểm danh với bộ lọc trạng thái và menu thao tác.")
    add_body(doc, "Để xem ai đã được ghi nhận trong một phiên, mở menu ba chấm ở cuối dòng và chọn Xem chi tiết. Danh sách hiển thị thiếu nhi, tài khoản, trạng thái và giờ ghi nhận ngay trên trang hiện tại.")
    add_screenshot(doc, "08-chi-tiet-phien-diem-danh.png", "Hình 10 - Popup danh sách tài khoản đã điểm danh.", "Popup hiển thị tài khoản thiếu nhi đã điểm danh trong phiên được chọn.")

    doc.add_page_break()
    doc.add_heading("3.3 Quản lý và tạo bài tập", level=2)
    add_body(doc, "Màn Bài tập gồm khu vực việc cần xử lý và danh mục bài tập. Giáo lý viên có thể tìm theo tên, lọc trạng thái, lọc lớp và tải lại danh sách.")
    add_screenshot(doc, "09-giao-ly-vien-bai-tap.png", "Hình 11 - Danh mục bài tập của Giáo lý viên.", "Trang bài tập giáo lý viên với hàng tìm kiếm, lọc trạng thái và lọc lớp.")
    add_body(doc, "Chọn Tạo bài tập để bắt đầu quy trình năm bước:")
    for text in (
        "Thông tin - nhập tên, mô tả, hình thức, thang điểm, điểm đạt và tệp đính kèm.",
        "Câu hỏi - thêm câu trắc nghiệm, trả lời ngắn hoặc tự luận; nhập đáp án và điểm.",
        "Người nhận - chọn một hoặc nhiều lớp, thời gian mở và hạn nộp.",
        "Thiết lập - cấu hình lượt làm, giới hạn thời gian, nộp trễ, công bố kết quả và trộn câu hỏi.",
        "Xem trước - kiểm tra toàn bộ đề trước khi Lưu bản nháp hoặc Phát hành bài tập.",
    ):
        add_list_item(doc, text, number_list, bold_prefix=text.split(" - ")[0] + " -")
    add_screenshot(doc, "10-tao-bai-tap.png", "Hình 12 - Bước Thông tin trong quy trình tạo bài tập.", "Màn hình tạo bài tập năm bước với biểu mẫu thông tin bài tập.")
    add_callout(doc, "Bản nháp", "Bản nháp chưa được gửi đến Thiếu nhi. Sau khi kiểm tra đủ câu hỏi, người nhận và thời gian, chọn Phát hành bài tập ở bước Xem trước.", tone="green")

    doc.add_heading("3.4 Chấm bài và thông báo lớp", level=2)
    for text in (
        "Mở Bài cần chấm để chọn bài có lượt nộp, xem từng bài làm và cập nhật điểm hoặc nhận xét cho phần tự luận.",
        "Mở Thông báo lớp để tạo thông báo cho lớp phụ trách. Kiểm tra đối tượng nhận và mức độ trước khi gửi.",
        "Các câu trắc nghiệm có đáp án được hệ thống xử lý tự động; phần tự luận cần giáo lý viên chấm trước khi công bố kết quả.",
    ):
        add_list_item(doc, text, bullet_list)

    doc.add_page_break()
    doc.add_heading("4. Hướng dẫn dành cho Thiếu nhi", level=1)
    doc.add_heading("4.1 Không gian học tập", level=2)
    add_body(doc, "Sau khi đăng nhập, Thiếu nhi được chuyển đến Lịch học. Thanh bên chỉ hiển thị các chức năng phục vụ việc học và điểm danh: Lịch học, Bài tập, Thông báo và Quét QR điểm danh. Trên điện thoại, chọn nút menu ở góc trên để mở thanh điều hướng.")

    doc.add_heading("4.2 Xem lịch học", level=2)
    for text in (
        "Mở Lịch học để xem lớp, phòng học và Giáo lý viên phụ trách.",
        "Khung Buổi học tiếp theo hiển thị ngày và giờ của buổi gần nhất.",
        "Trong lịch tuần, dùng nút mũi tên để chuyển tuần hoặc chọn Hôm nay để quay lại tuần hiện tại.",
        "Mỗi buổi học hiển thị khung giờ, tên lớp và phòng học. Nếu thông tin chưa đúng, liên hệ Giáo lý viên hoặc quản trị viên.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "11-thieu-nhi-lich-hoc.png", "Hình 13 - Lịch học tuần của Thiếu nhi.", "Màn hình Lịch học hiển thị buổi học tiếp theo, thông tin lớp và lịch theo từng ngày trong tuần.")

    doc.add_heading("4.3 Làm và nộp bài tập", level=2)
    for text in (
        "Mở Bài tập và chọn Cần làm để xem bài chưa hoàn thành; dùng Đã nộp hoặc Tất cả để xem các nhóm còn lại.",
        "Chọn Bắt đầu hoặc Tiếp tục, sau đó kiểm tra hạn nộp, thời gian làm, số lượt và số câu hỏi.",
        "Trả lời từng câu. Hệ thống tự động lưu tiến độ; có thể chọn Lưu ngay trước khi rời màn hình.",
        "Nếu bài tập cho phép, chọn Thêm tệp để đính kèm tệp bài làm theo định dạng và dung lượng hiển thị.",
        "Kiểm tra tiến độ rồi chọn Nộp bài. Sau khi nộp, bài xuất hiện trong Đã nộp và không thể sửa trừ khi Giáo lý viên mở lại.",
        "Khi Giáo lý viên công bố kết quả, mở lại bài để xem điểm và nhận xét.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "12-thieu-nhi-bai-tap.png", "Hình 14 - Danh sách bài tập của Thiếu nhi.", "Trang Bài tập với các tab Cần làm, Đã nộp và Tất cả; bài mới sẽ xuất hiện trong vùng danh sách.")

    doc.add_heading("4.4 Theo dõi thông báo", level=2)
    for text in (
        "Mở Thông báo để xem thông tin từ Giáo lý viên và các cập nhật liên quan đến bài tập.",
        "Dùng Tất cả để xem toàn bộ hoặc Chưa đọc để tập trung vào thông báo mới.",
        "Mở một thông báo để xem nội dung và đánh dấu đã đọc; chọn Đọc tất cả khi muốn xử lý toàn bộ thông báo chưa đọc.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "13-thieu-nhi-thong-bao.png", "Hình 15 - Màn hình Thông báo của Thiếu nhi.", "Màn hình Thông báo có bộ lọc Tất cả, Chưa đọc và hành động Đọc tất cả.")

    doc.add_heading("4.5 Điểm danh bằng QR", level=2)
    for text in (
        "Mở Quét QR điểm danh trên thiết bị sẽ dùng thường xuyên.",
        "Chọn Kích hoạt điện thoại này. Mỗi tài khoản chỉ nên gắn với thiết bị của chính mình.",
        "Khi giáo lý viên mở phiên, chọn Mở camera và đưa mã QR vào khung hình.",
        "Nếu camera không khả dụng, nhập đường dẫn hoặc mã do giáo lý viên cung cấp rồi chọn Điểm danh.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "14-thieu-nhi-qr.png", "Hình 16 - Màn hình kích hoạt và quét QR điểm danh.", "Màn hình điểm danh QR của Thiếu nhi với kích hoạt thiết bị, camera và nhập mã thủ công.")
    add_callout(doc, "Không chia sẻ", "Mỗi tài khoản chỉ được điểm danh một lần cho mỗi buổi học. Không gửi tài khoản hoặc mã phiên cho người khác.", tone="amber")

    doc.add_heading("4.6 Quản lý tài khoản", level=2)
    add_body(doc, "Chọn khu vực tài khoản ở cuối thanh bên để mở Tài khoản hoặc Đăng xuất. Trong trang Tài khoản, kiểm tra thông tin cá nhân và đổi mật khẩu khi cần. Luôn đăng xuất sau khi sử dụng thiết bị dùng chung.")

    doc.add_page_break()
    doc.add_heading("5. Hướng dẫn dành cho Phụ huynh", level=1)
    doc.add_heading("5.1 Theo dõi tổng quan", level=2)
    add_body(doc, "Không gian Phụ huynh tập trung vào việc theo dõi hồ sơ thiếu nhi đã liên kết, lịch học, lịch sử tham dự, bài tập, điểm thưởng và thông báo.")
    add_screenshot(doc, "14-phu-huynh-tong-quan.png", "Hình 17 - Không gian tổng quan của Phụ huynh.", "Trang tổng quan của tài khoản phụ huynh.")

    doc.add_heading("5.2 Xem hồ sơ thiếu nhi đã liên kết", level=2)
    for text in (
        "Mở Các con của tôi để xem danh sách hồ sơ đã được quản trị viên liên kết với gia đình.",
        "Kiểm tra mã thiếu nhi, tên thánh và trạng thái học tập.",
        "Nếu thiếu hồ sơ hoặc liên kết sai, liên hệ quản trị viên; phụ huynh không tự chuyển lớp hoặc gán hồ sơ.",
        "Thiếu nhi sử dụng tài khoản riêng để quét QR; phụ huynh không quét thay bằng tài khoản phụ huynh.",
    ):
        add_list_item(doc, text, number_list)
    add_screenshot(doc, "15-phu-huynh-thieu-nhi.png", "Hình 18 - Danh sách thiếu nhi liên kết với Phụ huynh.", "Trang Các con của tôi hiển thị hồ sơ thiếu nhi đã liên kết.")

    doc.add_page_break()
    doc.add_heading("6. Xử lý tình huống thường gặp", level=1)
    troubleshooting = [
        ("Không đăng nhập được", "Kiểm tra email, mật khẩu và trạng thái tài khoản. Nếu vẫn lỗi, liên hệ quản trị viên để xác nhận tài khoản chưa bị khóa."),
        ("Đăng ký xong nhưng chưa thấy lớp hoặc hồ sơ con", "Tự đăng ký chỉ tạo tài khoản và hồ sơ theo vai trò. Quản trị viên cần xếp lớp cho Thiếu nhi hoặc liên kết hồ sơ con với tài khoản Phụ huynh."),
        ("Không thấy chức năng", "Chức năng được ẩn theo vai trò và quyền. Đăng xuất, đăng nhập lại; nếu vẫn thiếu, đề nghị quản trị viên kiểm tra phân quyền."),
        ("Không thấy niên khóa, khối hoặc phòng khi sửa lớp", "Mở Danh mục lớp học, chọn đúng giáo xứ và kiểm tra danh mục đang ở trạng thái Đang sử dụng. Danh mục đã ngừng sử dụng không xuất hiện cho lựa chọn mới."),
        ("Không thấy lớp hoặc thiếu nhi", "Giáo lý viên chỉ thấy lớp được phân công. Quản trị viên cần kiểm tra phân công, niên khóa và trạng thái ghi danh."),
        ("Không lưu được bài tập", "Kiểm tra đủ tên bài, câu hỏi, điểm, người nhận và thời gian. Ở bước Xem trước, đọc thông báo lỗi và quay lại đúng bước để bổ sung."),
        ("Không quét được QR", "Kiểm tra thiết bị đã kích hoạt, quyền camera, phiên còn hiệu lực và tài khoản chưa điểm danh. Có thể dùng mã thủ công khi camera không khả dụng."),
        ("Dữ liệu chưa cập nhật", "Chọn nút tải lại nếu có. Tránh nhấn Lưu nhiều lần liên tiếp; chờ thông báo thành công trước khi rời trang."),
    ]
    table = doc.add_table(rows=1, cols=2)
    for index, text in enumerate(("Tình huống", "Cách xử lý")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_HEADER)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=10, color=NAVY, bold=True)
    for issue, resolution in troubleshooting:
        cells = table.add_row().cells
        for index, text in enumerate((issue, resolution)):
            run = cells[index].paragraphs[0].add_run(text)
            set_run_font(run, size=10, color=TEXT, bold=index == 0)
    repeat_table_header(table.rows[0])
    set_fixed_table_geometry(table, [2760, 6600])
    set_table_borders(table)

    doc.add_page_break()
    doc.add_heading("7. Phụ lục - tài khoản dữ liệu mẫu", level=1)
    add_callout(doc, "Chỉ dùng môi trường local", "Các tài khoản dưới đây phục vụ kiểm thử nội bộ. Khi triển khai thật, sử dụng tài khoản do quản trị viên tạo và tuyệt đối không dùng mật khẩu mẫu.", tone="amber")
    add_demo_accounts_table(doc)

    doc.add_page_break()
    doc.add_heading("Checklist trước khi kết thúc phiên làm việc", level=2)
    for text in (
        "Đã lưu hoặc phát hành dữ liệu cần thiết.",
        "Không còn biểu mẫu đang chỉnh sửa dở trên thiết bị dùng chung.",
        "Đã đóng camera QR nếu không còn sử dụng.",
        "Đã đăng xuất khỏi thiết bị công cộng hoặc thiết bị của người khác.",
    ):
        add_list_item(doc, text, bullet_list)

    doc.add_heading("Khi cần hỗ trợ", level=2)
    add_body(doc, "Gửi cho quản trị viên đủ thông tin để tái hiện lỗi; tránh chỉ mô tả chung là hệ thống không hoạt động.")
    for text in (
        "Vai trò và tài khoản đang sử dụng (không gửi mật khẩu).",
        "Tên trang, thời điểm xảy ra lỗi và thao tác vừa thực hiện.",
        "Ảnh màn hình có thông báo lỗi; che dữ liệu cá nhân nếu gửi qua kênh công cộng.",
    ):
        add_list_item(doc, text, bullet_list)
    add_callout(doc, "Gợi ý", "Thử tải lại trang một lần và đăng nhập lại trước khi báo lỗi. Nếu lỗi lặp lại, giữ nguyên màn hình để quản trị viên kiểm tra.")

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)
